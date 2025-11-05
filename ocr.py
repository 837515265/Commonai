import fitz
import paddle
import docx
from module.utils import to_list
import numpy as np
from PIL import Image
from loguru import logger
from pathlib import Path
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import tempfile, shutil, re
import tempfile, io
from collections import Counter
import subprocess, shlex, sys, os
### PATCH‑A：文件后缀白名单
SUPPORTED_EXTS = {'pdf', 'jpg', 'jpeg', 'png', 'docx', 'doc', 'txt'}
COMPRESS_EXTS = {'zip', 'rar'}
SUPPORTED_EXTS |= COMPRESS_EXTS

### PATCH‑B：安全解压函数（放在 ocr.py 里即可）
import zipfile
try:
    import rarfile           # pip install rarfile；系统需安装 unrar / rar
except ImportError:
    rarfile = None
from collections import Counter

def is_watermark_dominated(text: str, top_rate: float = 0.5) -> bool:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return True
    top_cnt = Counter(lines).most_common(1)[0][1]
    return top_cnt / len(lines) >= top_rate
def get_sevenz_path() -> Path:
    """
    返回项目内 7‑Zip 可执行文件的 Path。
    每次调用都会重新计算和 exists() 校验。
    """
    base_dir = Path(__file__).resolve().parent          # ocr.py 所在目录
    seven = base_dir / "tools" / ("7z.exe" if sys.platform.startswith("win") else "7zz")

    if not seven.exists():
        raise RuntimeError(f"7‑Zip executable not found at: {seven}")

    # Linux/macOS 确保有执行权限；Windows 不需要
    if not sys.platform.startswith("win"):
        mode = seven.stat().st_mode
        if not (mode & 0o111):
            seven.chmod(mode | 0o111)

    return seven


def safe_extract(src_path: Path, dst_dir: Path) -> list[Path]:
    """
    解压 zip / rar / 7z   → dst_dir
    返回解压出的、后缀在 SUPPORTED_EXTS – COMPRESS_EXTS 内的文件列表。

    依赖项目自带单文件 7‑Zip (tools/7z.exe 或 tools/7zz)；不用 rarfile，不走系统安装。
    """
    extracted: list[Path] = []
    suffix = src_path.suffix.lower()
    dst_dir.mkdir(parents=True, exist_ok=True)

    # ---------- ZIP 使用标准库 ----------
    if suffix == ".zip":
        with zipfile.ZipFile(src_path, "r") as zf:
            zf.extractall(dst_dir)

    # ---------- RAR / 7Z 使用项目内 7‑Zip ----------
    elif suffix in (".rar", ".7z"):
        SevenZ = str(get_sevenz_path())
        cmd = [SevenZ, "x", "-y", str(src_path), f"-o{dst_dir}"]  # x=解压  -y=自动Yes
        res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
        if res.returncode != 0:
            msg = res.stdout.decode(errors="ignore")
            raise RuntimeError(f"7‑Zip 解压失败 (code {res.returncode}):\n{msg}")

    # ---------- 其它类型 ----------
    else:
        # 如果不是压缩包，直接返回空列表
        return extracted

    # ---------- 过滤受支持的文件 ----------
    for p in dst_dir.rglob("*"):
        if p.is_file():
            ext = p.suffix.lstrip(".").lower()
            if ext in (SUPPORTED_EXTS - COMPRESS_EXTS):
                extracted.append(p)

    return extracted

# ------------------------------------------------------------------
class OcrError(Exception):
    def __init__(self, message):
        super().__init__(message)
        self.message = message

    def __str__(self):
        return f'OcrError:{self.message}'

class OcrTool():
    def __init__(self, pipeline, page_break_text='\n----- PAGE BREAK -----\n', raise_error=True):
        self.pipeline = pipeline
        self.page_break_text = page_break_text
        self.raise_error = raise_error

        log_name = 'OCR'
        logger.add(f"logs/module_{log_name}_{{time:YYYY-MM-DD}}.log",
               level="INFO",
               rotation="00:00",
               filter=lambda record: record["extra"].get("name") == log_name,
               enqueue=False,
               buffering=1)
        self.logger = logger.bind(name=log_name)

    # @staticmethod
    # def is_scanned_pdf(pdf_path):
    #     doc = fitz.open(pdf_path)
    #     for page in doc:
    #         text = page.get_text()
    #         if not text.strip():
    #             return True # 没有文本，则可能是扫描版
    #     return False  # 发现可选文本层，说明是电子版

    def read_pdf(self,file_path):
        try:
            doc = fitz.open(file_path)
            page_text_Dict = {}

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                # 清洗文本
                text = re.sub(r'[ \t\xa0]+', ' ', text)
                # 先统一换行符为 \n
                text = re.sub(r'\r\n|\r|\f', '\n', text)
                # 去除多余的空行，只保留单个换行
                text = re.sub(r'\n\s*\n', '\n', text)
                page_text_Dict[page_num + 1] = text

            doc.close()
            self.logger.info(f'fitz 原生提取完成: {file_path}')
            return page_text_Dict

        except Exception as e:
            self.logger.error(f'fitz 原生提取失败: {file_path}')
            self.logger.exception(e)
            if self.raise_error:
                raise OcrError('fitz 原生提取失败')
    def pdf_ocr(self, file_path, **kwargs):
        try:
            pdf_file = to_list(file_path)
            pdf_file = [Path(file).as_posix() for file in pdf_file]

            paras = dict(text_det_box_thresh=0.4,
                                      text_det_unclip_ratio=1.5,
                                     use_textline_orientation=False)
            paras.update(kwargs)

            output = self.pipeline.predict(pdf_file, **paras)

            page_text_Dict = {}
            for index, page in enumerate(output):
                page_text = '\n'.join(page['rec_texts'])

                # 清洗逻辑：统一空白和换行
                page_text = re.sub(r'[ \t\xa0]+', ' ', page_text)  # 统一空格类字符为一个空格
                page_text = re.sub(r'\r\n|\r|\f', '\n', page_text)  # 统一换行符为 \n
                page_text = re.sub(r'(\n\s*)+', '\n\n', page_text)  # 多个换行/空行合并为段落间空行

                page_num = index + 1
                page_text_Dict[page_num] = page_text.strip()  # 去头尾空白
            doc_text = self.page_break_text.join(page_text_Dict[page_num] for page_num in sorted(page_text_Dict.keys()))
            self.logger.info(f'pdf ocr成功:{file_path}')
            return doc_text, page_text_Dict
        except Exception as e:
            self.logger.error(f'pdf ocr失败:{file_path}')
            self.logger.exception(e)
            if self.raise_error:
                raise OcrError(f'pdf ocr失败')
        finally:
            paddle.device.cuda.empty_cache()

    def image_ocr(self, file_input, **kwargs):
        try:
            paras = dict(
                text_det_box_thresh=0.4,
                text_det_unclip_ratio=1.5,
                use_textline_orientation=False
            )
            paras.update(kwargs)

            page_text_Dict = {}

            if isinstance(file_input, Image.Image):
                # 单张 PIL 图像
                img = np.array(file_input.convert("RGB"))[:, :, ::-1]
                output = self.pipeline.predict(img, **paras)
                page_text = '\n'.join(output['rec_texts'])
                page_text_Dict[1] = re.sub(r'\s+', '\n', page_text)
                doc_text = page_text_Dict[1]
                return doc_text, page_text_Dict

            # 如果是文件路径或路径列表
            image_file = to_list(file_input)
            image_file = [Path(file).as_posix() for file in image_file]
            images = [np.array(Image.open(file).convert('RGB'))[:, :, ::-1] for file in image_file]

            output = self.pipeline.predict(images, **paras)

            for index, page in enumerate(output):
                page_text = '\n'.join(page['rec_texts'])
                page_num = index + 1
                page_text_Dict[page_num] = re.sub('\s+', '\n', page_text)

            doc_text = self.page_break_text.join(
                page_text_Dict[page_num] for page_num in sorted(page_text_Dict.keys())
            )

            self.logger.info(f'image ocr成功:{file_input}')
            return doc_text, page_text_Dict

        except Exception as e:
            self.logger.error(f'image ocr失败:{file_input}')
            self.logger.exception(e)
            raise OcrError(f'image ocr失败')
        finally:
            paddle.device.cuda.empty_cache()

    def read_txt(self, file_path):
        try:
            file_path = to_list(file_path)
            text_List = []
            for file in file_path:
                with open(file, encoding='utf-8', errors='ignore') as f:
                    text_List.append(f.read())
            doc_text = self.page_break_text.join(text_List)

            page_text_List = doc_text.split(self.page_break_text)
            page_text_Dict = dict(zip(range(1, len(page_text_List) + 1), page_text_List))
            self.logger.info(f'txt解析成功:{file_path}')
            return doc_text, page_text_Dict
        except Exception as e:
            self.logger.error(f'txt解析失败:{file_path}')
            self.logger.exception(e)
            if self.raise_error:
                raise OcrError(f'txt解析失败')

    # def read_docx(self, file_path):
    #     try:
    #         doc = docx.Document(file_path)
    #         doc_text = self.page_break_text.join([para.text for para in doc.paragraphs])
    #         self.logger.info(f'{len(doc.paragraphs)=}')
    #         page_text_Dict = dict(zip(range(1,len(doc.paragraphs) + 1), [para.text for para in doc.paragraphs]))
    #         self.logger.info(f'docx解析成功:{file_path}')
    #         return doc_text, page_text_Dict
    #     except Exception as e:
    #         self.logger.error(f'docx解析失败:{file_path}')
    #         self.logger.exception(e)
    #         if self.raise_error:
    #             raise OcrError(f'docx解析失败')
    def read_docx(self, file_path: Path, text_threshold: int = 20):
        """
        • 若文字 > text_threshold，则直接返回段落文本
        • 否则尝试提取所有嵌入图片，调用 image_ocr
        • 提取失败或无图时返回原文本
        """
        doc = docx.Document(file_path)

        # 1) 聚合段落文本
        para_texts = [p.text for p in doc.paragraphs]
        plain_txt = "\n".join(para_texts).strip()

        if len(plain_txt) > text_threshold:
            doc_text = self.page_break_text.join(para_texts)
            return doc_text, {1: doc_text}

        # 2) 提取图片
        tmp_dir = Path(tempfile.mkdtemp(prefix="docx_img_"))
        image_files: list[Path] = []
        try:
            for rel in doc.part.rels.values():
                if rel.reltype == RT.IMAGE:  # ★ 关键：官方常量判断
                    img_bytes = rel.target_part.blob
                    name = Path(rel.target_ref).name or "image"
                    out = tmp_dir / name
                    idx = 1
                    while out.exists():
                        out = tmp_dir / f"{out.stem}_{idx}{out.suffix}"
                        idx += 1
                    with open(out, "wb") as f:
                        f.write(img_bytes)
                    image_files.append(out)

            if not image_files:
                # 无图仍返回原文本
                doc_text = self.page_break_text.join(para_texts)
                return doc_text, {1: doc_text}

            # 3) OCR
            ocr_text, page_dict = self.image_ocr(image_files)
            return ocr_text, page_dict

        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)
    def ocr_files(self, file_path_List):
        file_path_List = to_list(file_path_List)
        file_path_List = [Path(file_path) for file_path in file_path_List]
        self.logger.info(f'{file_path_List=}')

        extra_files: list[Path] = []
        for p in list(file_path_List):  # list() 复制，避免遍历时改原列表
            if p.suffix.lstrip('.').lower() in COMPRESS_EXTS:
                unzip_dir = p.parent / f"{p.stem}_unzip"
                unzip_dir.mkdir(exist_ok=True, parents=True)
                try:
                    inner = safe_extract(p, unzip_dir)
                    extra_files.extend(inner)
                    self.logger.info(f'解压 {p.name} -> {len(inner)} 个文件')
                except Exception as e:
                    self.logger.error(f'解压 {p.name} 失败: {e}')

        # 把解压得到的文件追加到待处理列表
        file_path_List.extend(extra_files)


        all_text = ''
        for index, file_path in enumerate(file_path_List):
            suffix = file_path.suffix.lstrip('.').lower()
            self.logger.info(f'{suffix=}')
            doc_text = ''
            page_text_Dict = {}
            if suffix == 'pdf':
                try:
                    doc = fitz.open(file_path)

                    pages_to_check = min(3, len(doc))
                    merged_txt = ""
                    eff_len = 0
                    for pg in range(pages_to_check):
                        pg_txt = doc[pg].get_text("text")
                        merged_txt += pg_txt + "\n"
                        eff_len += count_effective_chars(pg_txt)

                    watermark_heavy = is_watermark_dominated(merged_txt, top_rate=0.5)

                    # 同时满足「有效字符少」或「水印占比高」就 OCR
                    use_ocr_for_all = (eff_len < 60) or watermark_heavy
                    # ---------------------------

                    # first_page = doc[0]
                    # text = re.sub(r'[ \t\xa0]+', ' ', first_page.get_text("text"))
                    # text = re.sub(r'\n{2,}', '\n', text).strip()
                    # use_ocr_for_all = len(text) < 10

                    if use_ocr_for_all:
                        self.logger.info(f"第 1 页字符太少，使用 OCR 处理整个文件")
                        doc_text, page_text_Dict = self.pdf_ocr(file_path)
                    else:
                        self.logger.info(f"第 1 页文本正常，使用 fitz 原生提取整个文件")
                        page_text_Dict = self.read_pdf(file_path)
                        doc_text = self.page_break_text.join(
                            page_text_Dict[p] for p in sorted(page_text_Dict.keys())
                        )

                    doc.close()
                    doc_text = self.page_break_text.join(
                        page_text_Dict[page_num] for page_num in sorted(page_text_Dict.keys()))
                    self.logger.info(f'pdf逐页混合处理完成:{file_path}')
                except Exception as e:
                    self.logger.error(f'pdf处理失败:{file_path}')
                    self.logger.exception(e)
                    if self.raise_error:
                        raise OcrError(f'pdf处理失败')
            elif suffix in ('jpg','jpeg','png'):
                doc_text, page_text_Dict = self.image_ocr(file_path)
            elif suffix == 'txt':
                doc_text, page_text_Dict = self.read_txt(file_path)
            elif suffix  in ('docx', 'doc'):
                doc_text, page_text_Dict = self.read_docx(file_path)
            else:
                self.logger.error(f'{file_path}的类型无法处理')

            if all_text:
                all_text += self.page_break_text + doc_text
            else:
                all_text = doc_text
        page_text_List = all_text.split(self.page_break_text)
        page_text_Dict = dict(zip(range(1, len(page_text_List) + 1), page_text_List))
        return all_text, page_text_Dict

    def get_duplicate_line(self, lines, most_common=30, min_length=5, min_count=10):
        line_counter = Counter(lines)
        duplicate_line_List = [line for line,count in line_counter.most_common(most_common) if len(line) >= min_length and count >= min_count
                               ]

        return [line for line in duplicate_line_List if line.strip() != self.page_break_text.strip()]

    def remove_duplicate_sentences(self, doc_text, page_text_Dict, **kwargs):
        lines = doc_text.split('\n')
        duplicate_line_List = self.get_duplicate_line(lines, **kwargs)
        lines = [line for line in lines if line not in duplicate_line_List]
        doc_text = '\n'.join(lines)
        for page_num in page_text_Dict:
            page_text_Dict[page_num] = '\n'.join([line for line in page_text_Dict[page_num].split('\n') if line not in duplicate_line_List])
        return doc_text, page_text_Dict
def count_effective_chars(txt: str) -> int:
    """
    仅保留中文、英文、数字作为“有效字符”
    """
    txt = re.sub(r'[\s\r\n\t\xa0]+', '', txt)           # 去空白
    txt = re.sub(r'[^\w\u4e00-\u9fa5]', '', txt)        # 去符号（保留字母数字中文）
    return len(txt)